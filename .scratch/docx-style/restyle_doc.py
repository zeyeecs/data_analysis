from copy import deepcopy
import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Cm, Pt, RGBColor


SOURCE = Path("/Users/lanse/Documents/sjkx/.scratch/docx-style/original.docx")
OUTPUT = Path("/Users/lanse/Documents/sjkx/.scratch/docx-style/成都启钥网络科技有限公司_样式优化版.docx")


URL_PATTERN = re.compile(r"(https?://\S+)")
SECTION_HEADING = re.compile(r"^[一二三四五六七八九十]+、")
NUMERIC_HEADING = re.compile(r"^\d+\.\s")
PLATFORM_OR_LABEL = re.compile(
    r"^(微信广告平台|抖音 / 巨量引擎|3D 游戏试玩案例|2D 游戏试玩案例|创意策划（黄金30秒体验设计）\s*：?)$"
)
FLOW_LINE = "需求分析 ➔ 创意策划 ➔ 技术开发 ➔ 测试优化 ➔ 平台对接"
BULLET_PREFIX = (
    "提升用户质量：",
    "降低获客成本：",
    "优化用户体验：",
    "增强品牌认知：",
    "引擎全面支持：",
    "极致包体控制：",
    "秒级加载体验：",
    "卓越性能表现：",
    "全链路数据追踪：",
    "优势：",
    "适合品类：",
    "转化率（CR）大幅增益：",
    "用户质量显著提高：",
    "次日留存率提升",
    "7日留存率提升",
    "付费转化率提升",
    "用户生命周期价值（LTV）提升",
    "获客成本深度优化：",
    "安装成本（CPI）降低",
    "有效用户获取成本降低",
    "总体 ROI 平均提升",
    "0-5秒：",
    "5-15秒：",
    "15-25秒：",
    "25-30秒：",
    "测试优化与平台对接：",
    "转化率提升保证：",
    "技术稳定性：",
    "数据透明度：",
)


def set_east_asia_font(run, font_name: str) -> None:
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)


def format_run(run, *, font_name: str, size: int, bold=False, color=None, italic=False) -> None:
    run.font.name = font_name
    set_east_asia_font(run, font_name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1F5AA6")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Aptos")
    fonts.set(qn("w:hAnsi"), "Aptos")
    fonts.set(qn("w:eastAsia"), "等线")
    r_pr.append(fonts)

    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "20")
    r_pr.append(size)

    size_cs = OxmlElement("w:szCs")
    size_cs.set(qn("w:val"), "20")
    r_pr.append(size_cs)

    new_run.append(r_pr)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def clone_style_font(dst_style, src_style) -> None:
    dst_style.font.name = src_style.font.name
    dst_style.font.size = src_style.font.size
    dst_style.font.bold = src_style.font.bold
    dst_style.font.italic = src_style.font.italic
    if src_style.font.color.rgb:
        dst_style.font.color.rgb = deepcopy(src_style.font.color.rgb)


def set_style_east_asia(style, font_name: str) -> None:
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)


def ensure_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Aptos"
    set_style_east_asia(normal, "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.space_before = Pt(0)

    if "Title" not in styles:
        title = styles.add_style("Title", WD_STYLE_TYPE.PARAGRAPH)
    else:
        title = styles["Title"]
    title.font.name = "Aptos"
    set_style_east_asia(title, "等线")
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string("17365D")
    title.paragraph_format.space_after = Pt(18)
    title.paragraph_format.space_before = Pt(0)

    h1 = styles["Heading 1"]
    h1.font.name = "Aptos"
    set_style_east_asia(h1, "等线")
    h1.font.size = Pt(15)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string("1F4E79")
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)

    h2 = styles["Heading 2"]
    h2.font.name = "Aptos"
    set_style_east_asia(h2, "等线")
    h2.font.size = Pt(12.5)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string("2F75B5")
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(6)

    h3 = styles["Heading 3"]
    h3.font.name = "Aptos"
    set_style_east_asia(h3, "等线")
    h3.font.size = Pt(11.5)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor.from_string("3F3F3F")
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)

    if "Body Bullet" not in styles:
        bullet = styles.add_style("Body Bullet", WD_STYLE_TYPE.PARAGRAPH)
    else:
        bullet = styles["Body Bullet"]
    clone_style_font(bullet, normal)
    set_style_east_asia(bullet, "宋体")
    bullet.paragraph_format.left_indent = Cm(0.75)
    bullet.paragraph_format.first_line_indent = Cm(-0.5)
    bullet.paragraph_format.space_after = Pt(6)
    bullet.paragraph_format.line_spacing = 1.3

    if "Case Link" not in styles:
        link_style = styles.add_style("Case Link", WD_STYLE_TYPE.PARAGRAPH)
    else:
        link_style = styles["Case Link"]
    clone_style_font(link_style, normal)
    set_style_east_asia(link_style, "宋体")
    link_style.paragraph_format.left_indent = Cm(0.75)
    link_style.paragraph_format.space_after = Pt(4)
    link_style.paragraph_format.line_spacing = 1.2

    if "Flow Line" not in styles:
        flow = styles.add_style("Flow Line", WD_STYLE_TYPE.PARAGRAPH)
    else:
        flow = styles["Flow Line"]
    flow.font.name = "Aptos"
    set_style_east_asia(flow, "等线")
    flow.font.size = Pt(11.5)
    flow.font.bold = True
    flow.font.color.rgb = RGBColor.from_string("1F4E79")
    flow.paragraph_format.space_before = Pt(6)
    flow.paragraph_format.space_after = Pt(10)
    flow.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def clear_paragraph_runs(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            paragraph._p.remove(child)


def replace_with_hyperlink(paragraph, prefix: str, url: str) -> None:
    clear_paragraph_runs(paragraph)
    run = paragraph.add_run(prefix)
    format_run(run, font_name="宋体", size=10.5)
    add_hyperlink(paragraph, url, url)


def configure_sections(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.4)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.header_distance = Cm(1.2)
        section.footer_distance = Cm(1.2)

        footer = section.footer
        if footer.paragraphs:
            p = footer.paragraphs[0]
        else:
            p = footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.text = ""
        run = p.add_run("成都启钥网络科技有限公司  ")
        format_run(run, font_name="等线", size=9, color="808080")
        run.add_break()
        page_run = p.add_run("第 ")
        format_run(page_run, font_name="等线", size=9, color="808080")
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        page_run._r.append(fld_begin)
        page_run._r.append(instr)
        page_run._r.append(fld_sep)
        page_text = OxmlElement("w:t")
        page_text.text = "1"
        page_run._r.append(page_text)
        page_run._r.append(fld_end)
        tail = p.add_run(" 页")
        format_run(tail, font_name="等线", size=9, color="808080")


def clean_empty_paragraphs(doc: Document) -> None:
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip():
            continue
        p = paragraph._element
        parent = p.getparent()
        if parent is not None:
            parent.remove(p)


def restyle_paragraph(paragraph, idx: int) -> None:
    text = paragraph.text.strip()
    if not text:
        return

    if idx == 0:
        paragraph.style = "Title"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            format_run(run, font_name="等线", size=22, bold=True, color="17365D")
        return

    if SECTION_HEADING.match(text):
        paragraph.style = "Heading 1"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            format_run(run, font_name="等线", size=15, bold=True, color="1F4E79")
        return

    if NUMERIC_HEADING.match(text) or text.endswith("案例") or text == "成功指标保障：":
        paragraph.style = "Heading 2"
        for run in paragraph.runs:
            format_run(run, font_name="等线", size=12.5, bold=True, color="2F75B5")
        return

    if PLATFORM_OR_LABEL.match(text):
        paragraph.style = "Heading 3"
        for run in paragraph.runs:
            format_run(run, font_name="等线", size=11.5, bold=True, color="3F3F3F")
        return

    if text == FLOW_LINE:
        paragraph.style = "Flow Line"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            format_run(run, font_name="等线", size=11.5, bold=True, color="1F4E79")
        return

    if text.startswith("案例") and "http" in text:
        paragraph.style = "Case Link"
        match = URL_PATTERN.search(text)
        if match:
            url = match.group(1)
            prefix = text[: match.start(1)]
            replace_with_hyperlink(paragraph, prefix, url)
        return

    if text.startswith("官方网站：") and "http" in text:
        paragraph.style = "Case Link"
        match = URL_PATTERN.search(text)
        if match:
            url = match.group(1)
            prefix = text[: match.start(1)]
            replace_with_hyperlink(paragraph, prefix, url)
        return

    if text.startswith(BULLET_PREFIX):
        paragraph.style = "Body Bullet"
        for run in paragraph.runs:
            format_run(run, font_name="宋体", size=10.5)
        if paragraph.runs:
            paragraph.runs[0].bold = True
        return

    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in paragraph.runs:
        format_run(run, font_name="宋体", size=10.5)


def main() -> None:
    doc = Document(SOURCE)
    ensure_styles(doc)
    clean_empty_paragraphs(doc)

    for idx, paragraph in enumerate(doc.paragraphs):
        restyle_paragraph(paragraph, idx)

    configure_sections(doc)
    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
